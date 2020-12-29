from docx import Document
from docx.shared import Inches

# make rows bold
def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

# creating the bill for print
def createBill(data, bill_no):
    # cheking the bill state
    if data['payment']+data['amountToBe']-data['ItemTotal'] == 0:
        state = False
    else:
        state = True

    document = Document()

    p = document.add_paragraph()
    p.add_run('\n\n\n\n\n\n\n\n\n')

    document.add_heading('Bill No: {}'.format(data['bill_no']))

    # sorting the rented item list
    data['RentedItems'] = sorted(data['RentedItems'], key=lambda item: item[1])

    table = document.add_table(rows=4, cols=2)
    table.rows[0].cells[0].text = 'දිනය'
    table.rows[0].cells[1].text = data['date']
    table.rows[1].cells[0].text = 'ජා. අංකය'
    table.rows[1].cells[1].text = data['id']
    table.rows[2].cells[0].text = 'නම'
    table.rows[2].cells[1].text = data['name']
    table.rows[3].cells[0].text = 'දු. අංකය'
    table.rows[3].cells[1].text = data['tp']

    # rented item detailed section
    document.add_heading('Rented Item details')

    if state:
        table = document.add_table(rows=1, cols=6)
        table.style = 'Table Grid'  
        table.rows[0].cells[0].text = 'භාණ්ඩ වර්ගය'
        table.rows[0].cells[1].text = 'කුලියට ගත් දිනය'
        table.rows[0].cells[2].text = 'දින ගණන'
        table.rows[0].cells[3].text = 'ප්‍රමාණය'
        table.rows[0].cells[4].text = 'දිනකට කුලිය'
        table.rows[0].cells[5].text = 'මුදල'

        make_rows_bold(table.rows[0])

        for item in data['RentedItems']:
            row_cells = table.add_row().cells
            row_cells[0].text = item[0]
            row_cells[1].text = item[1]
            row_cells[2].text = item[2]
            row_cells[3].text = item[3]
            row_cells[4].text = item[4]
            row_cells[5].text = item[5]

        # total Row
        row = table.add_row()
        row_cells = row.cells
        row_cells[1].text = "Total"
        row_cells[5].text = '{:10,.2f}'.format(data['ItemTotal'])
        make_rows_bold(row)
    else:
        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'  
        table.rows[0].cells[0].text = 'භාණ්ඩ වර්ගය'
        table.rows[0].cells[1].text = 'කුලියට ගත් දිනය'
        table.rows[0].cells[2].text = 'ප්‍රමාණය'
        table.rows[0].cells[3].text = 'දිනකට කුලිය'

        make_rows_bold(table.rows[0])

        for item in data['RentedItems']:
            row_cells = table.add_row().cells
            row_cells[0].text = item[0]
            row_cells[1].text = item[1]
            row_cells[2].text = item[3]
            row_cells[3].text = item[4]

    document.add_heading('Payment Details', level=1)
    
    if state:
        table = document.add_table(rows=5, cols=2)
        table.style = 'Table Grid'  
        table.rows[0].cells[0].text = 'හිඟ මුදල'
        table.rows[0].cells[1].text = 'Rs. {:10,.2f}'.format(data['payment']+data['amountToBe']-data['ItemTotal'])
        table.rows[1].cells[0].text = 'මෙම බිල්පතෙහි වටිනාකම'
        table.rows[1].cells[1].text = 'Rs. {:10,.2f}'.format(data['ItemTotal'])
        table.rows[2].cells[0].text = 'මුළු මුදල'
        table.rows[2].cells[1].text = '            Rs. {:10,.2f}'.format(data['payment']+data['amountToBe'])
        table.rows[3].cells[0].text = 'අද දින ගෙවන ලද මුදල'
        table.rows[3].cells[1].text = '-(Rs. {:10,.2f})'.format(data['payment'])
        table.rows[4].cells[0].text = 'හිඟ මුදල'
        table.rows[4].cells[1].text = '            Rs. {:10,.2f}'.format(data['amountToBe'])
        make_rows_bold(table.rows[0], table.rows[1], table.rows[2], table.rows[3], table.rows[4])
    else:
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'  
        table.rows[0].cells[0].text = 'අද දින ගෙවන ලද මුදල'
        table.rows[0].cells[1].text = 'Rs. {:10,.2f}'.format(data['payment'])
        make_rows_bold(table.rows[0])
    

    document.save('docs/{}.docx'.format(bill_no))


data = {
    'bill_no': '0001',
    'id': '199732400730',
    'name': 'Dinindu Udana Thilakarathna',
    'tp': '0777186434',
    'date': '2020-12-10 17:36:41',
    'payment': 13000,
    'amountToBe': 12000,
    'RentedItems': [
        ['Poker sdrfgsdfgsdfg sdfg sdfgs', '2020-12-10', '2', '2', '1,500', '12,000'],
        ['Poker', '2020-12-10', '2', '2', '1,500', '12,000'],
        ['Poker', '2020-12-10', '2', '2', '1,500', '12,000'],
        ['Poker', '2020-12-10', '2', '2', '1,500', '12,000'],
    ],
    'ItemTotal': 10000,
}

createBill(data, '0012')